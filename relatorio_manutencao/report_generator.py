import io
import re
import zipfile
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import (
    WD_CELL_VERTICAL_ALIGNMENT,
    WD_TABLE_ALIGNMENT,
)
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageOps


# ============================================================
# CONFIGURAÇÕES
# ============================================================

AZUL = "1F4E78"
VERMELHO = "FF0000"

EXTENSOES = {
    ".jpg",
    ".jpeg",
    ".png",
    ".webp",
    ".heic",
}

# Ajuste para o caminho real dentro do seu projeto.
CAMINHO_LOGO = Path(__file__).resolve().parent.parent / "modelos" / "logo_rdo.png"


# ============================================================
# FUNÇÕES AUXILIARES
# ============================================================

def _normalizar(texto: str) -> str:
    return re.sub(
        r"[^A-Z0-9]+",
        " ",
        (texto or "").upper()
    ).strip()


def _categoria(tipo_ativo: str) -> str:
    texto = _normalizar(tipo_ativo)

    if "DISJUNTOR" in texto:
        return "DISJUNTOR"

    if "SECCION" in texto:
        return "SECCIONADORA"

    if "GERADOR" in texto or "GMG" in texto:
        return "GMG"

    return "GENERICO"


def _sombrear(celula, cor: str) -> None:
    propriedades = celula._tc.get_or_add_tcPr()

    elemento = propriedades.find(qn("w:shd"))

    if elemento is None:
        elemento = OxmlElement("w:shd")
        propriedades.append(elemento)

    elemento.set(qn("w:fill"), cor)


def _set_cell_margins(
    cell,
    top=80,
    start=80,
    bottom=80,
    end=80,
):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcMar = tcPr.first_child_found_in("w:tcMar")

    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)

    for m, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tcMar.find(qn(f"w:{m}"))

        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)

        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _texto_celula(
    celula,
    texto: str,
    rotulo: bool = False,
    tamanho=9,
) -> None:

    celula.text = ""

    paragrafo = celula.paragraphs[0]
    paragrafo.paragraph_format.space_before = Pt(0)
    paragrafo.paragraph_format.space_after = Pt(0)

    run = paragrafo.add_run(str(texto or ""))

    run.font.name = "Arial"
    run.font.size = Pt(tamanho)
    run.bold = rotulo

    if rotulo:
        run.font.color.rgb = RGBColor(
            255,
            255,
            255
        )

        _sombrear(
            celula,
            AZUL
        )

    celula.vertical_alignment = (
        WD_CELL_VERTICAL_ALIGNMENT.CENTER
    )

    _set_cell_margins(celula)


# ============================================================
# CONFIGURAÇÃO DO DOCUMENTO
# ============================================================

def _configurar(doc: Document) -> None:

    secao = doc.sections[0]

    # A4
    secao.page_width = Cm(21)
    secao.page_height = Cm(29.7)

    # Aproximação do relatório original
    secao.left_margin = Cm(1.8)
    secao.right_margin = Cm(1.8)
    secao.top_margin = Cm(1.7)
    secao.bottom_margin = Cm(1.5)

    normal = doc.styles["Normal"]

    normal.font.name = "Arial"
    normal.font.size = Pt(10)

    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.0

    # ----------------------------------
    # Heading 1
    # ----------------------------------

    h1 = doc.styles["Heading 1"]

    h1.font.name = "Arial"
    h1.font.size = Pt(11)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)

    h1.paragraph_format.space_before = Pt(8)
    h1.paragraph_format.space_after = Pt(5)

    # ----------------------------------
    # Heading 2
    # ----------------------------------

    h2 = doc.styles["Heading 2"]

    h2.font.name = "Arial"
    h2.font.size = Pt(10)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)

    h2.paragraph_format.space_before = Pt(6)
    h2.paragraph_format.space_after = Pt(4)


# ============================================================
# CAMPOS WORD
# ============================================================

def _campo_word(paragrafo, instrucao: str) -> None:
    """
    Insere um campo do Word.

    Exemplos:

    PAGE

    TOC \\o "1-2" \\h \\z \\u
    """

    run = paragrafo.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(
        qn("w:fldCharType"),
        "begin"
    )

    instr = OxmlElement("w:instrText")
    instr.set(
        qn("xml:space"),
        "preserve"
    )

    instr.text = instrucao

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(
        qn("w:fldCharType"),
        "separate"
    )

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(
        qn("w:fldCharType"),
        "end"
    )

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(fld_end)


def _atualizar_campos_ao_abrir(doc: Document) -> None:

    settings = doc.settings._element

    update_fields = settings.find(
        qn("w:updateFields")
    )

    if update_fields is None:
        update_fields = OxmlElement(
            "w:updateFields"
        )

        settings.append(update_fields)

    update_fields.set(
        qn("w:val"),
        "true"
    )


def _iniciar_numeracao(secao, numero=1):

    sectPr = secao._sectPr

    pgNumType = sectPr.find(
        qn("w:pgNumType")
    )

    if pgNumType is None:
        pgNumType = OxmlElement(
            "w:pgNumType"
        )

        sectPr.append(pgNumType)

    pgNumType.set(
        qn("w:start"),
        str(numero)
    )


# ============================================================
# CAPA
# ============================================================

def _adicionar_logo(
    doc: Document,
    caminho_logo: Path | None,
):

    if not caminho_logo:
        return

    if not Path(caminho_logo).exists():
        return

    p = doc.add_paragraph()

    p.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    p.paragraph_format.space_after = Pt(5)

    p.add_run().add_picture(
        str(caminho_logo),
        width=Cm(6.5),
    )


def _quadro_titulo_capa(
    doc: Document,
    periodicidade: str,
    tipo_ativo: str,
    subestacao: str,
):

    # Espaço semelhante ao documento original
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    tabela = doc.add_table(
        rows=1,
        cols=1,
    )

    tabela.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    tabela.style = "Table Grid"

    tabela.autofit = False

    celula = tabela.cell(0, 0)
    celula.width = Cm(17.5)

    _set_cell_margins(
        celula,
        top=350,
        bottom=350,
        start=150,
        end=150,
    )

    # --------------------------------------
    # Título
    # --------------------------------------

    p = celula.paragraphs[0]

    p.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    p.paragraph_format.space_after = Pt(5)

    r = p.add_run(
        f"RELATÓRIO {periodicidade.upper()} "
        f"{tipo_ativo.upper()}"
    )

    r.font.name = "Arial"
    r.font.size = Pt(17)
    r.bold = True

    # --------------------------------------
    # Manutenção
    # --------------------------------------

    p = celula.add_paragraph()

    p.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    p.paragraph_format.space_after = Pt(3)

    r = p.add_run(
        f"MANUTENÇÃO PREVENTIVA "
        f"{periodicidade.upper()}"
    )

    r.font.name = "Arial"
    r.font.size = Pt(13)

    # --------------------------------------
    # Empresa
    # --------------------------------------

    p = celula.add_paragraph()

    p.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    p.paragraph_format.space_after = Pt(3)

    r = p.add_run(
        "RIALMA TRANSMISSORA "
        "DE ENERGIA V S.A."
    )

    r.font.name = "Arial"
    r.font.size = Pt(12)

    # --------------------------------------
    # Subestação
    # --------------------------------------

    p = celula.add_paragraph()

    p.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    r = p.add_run(
        "SUBESTAÇÃO "
    )

    r.font.name = "Arial"
    r.font.size = Pt(13)

    r = p.add_run(
        subestacao.upper()
    )

    r.font.name = "Arial"
    r.font.size = Pt(13)

    # Igual ao relatório de referência
    r.font.color.rgb = RGBColor(
        255,
        0,
        0
    )


def _quadro_corpo_tecnico(
    doc: Document,
    corpo_tecnico: list[dict],
    numero_os: str,
    numero_apr: str,
    periodo: str,
    concessao: str,
):

    # Empurra o quadro para parte inferior
    for _ in range(5):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    qtd = max(
        len(corpo_tecnico),
        1
    )

    tabela = doc.add_table(
        rows=qtd + 5,
        cols=2,
    )

    tabela.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    tabela.style = "Table Grid"

    tabela.autofit = False

    tabela.columns[0].width = Cm(10.5)
    tabela.columns[1].width = Cm(6.0)

    # ======================================
    # CORPO TÉCNICO
    # ======================================

    cab = tabela.cell(0, 0).merge(
        tabela.cell(0, 1)
    )

    cab.text = ""

    p = cab.paragraphs[0]

    p.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    r = p.add_run(
        "CORPO TÉCNICO"
    )

    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(10)

    # ======================================
    # PROFISSIONAIS
    # ======================================

    if corpo_tecnico:

        for i, pessoa in enumerate(
            corpo_tecnico,
            start=1
        ):

            nome = pessoa.get(
                "nome",
                ""
            )

            funcao = pessoa.get(
                "funcao",
                ""
            )

            c1 = tabela.cell(i, 0)
            c2 = tabela.cell(i, 1)

            c1.text = ""
            c2.text = ""

            p1 = c1.paragraphs[0]

            r = p1.add_run(
                "Nome: "
            )

            r.bold = True
            r.font.name = "Arial"
            r.font.size = Pt(9)

            r = p1.add_run(nome)

            r.font.name = "Arial"
            r.font.size = Pt(9)

            p2 = c2.paragraphs[0]

            r = p2.add_run(
                "Função: "
            )

            r.bold = True
            r.font.name = "Arial"
            r.font.size = Pt(9)

            r = p2.add_run(funcao)

            r.font.name = "Arial"
            r.font.size = Pt(9)

    else:

        c1 = tabela.cell(1, 0)
        c2 = tabela.cell(1, 1)

        c1.text = "Nome:"
        c2.text = "Função:"

    # ======================================
    # DADOS ADICIONAIS
    # ======================================

    primeira_linha = qtd + 1

    dados = [
        (
            "OS",
            numero_os
        ),
        (
            "APR",
            numero_apr
        ),
        (
            "Período",
            periodo
        ),
        (
            "Concessão",
            concessao
        ),
    ]

    for offset, (rotulo, valor) in enumerate(dados):

        linha = primeira_linha + offset

        celula = tabela.cell(
            linha,
            0
        ).merge(
            tabela.cell(
                linha,
                1
            )
        )

        celula.text = ""

        p = celula.paragraphs[0]

        r = p.add_run(
            f"{rotulo}: "
        )

        r.bold = True
        r.font.name = "Arial"
        r.font.size = Pt(9)

        r = p.add_run(
            str(valor or "")
        )

        r.font.name = "Arial"
        r.font.size = Pt(9)


def _capa(
    doc: Document,
    dados: dict,
    corpo_tecnico: list[dict],
    numero_os: str,
    numero_apr: str,
    concessao: str,
    caminho_logo: Path | None,
):

    _adicionar_logo(
        doc,
        caminho_logo
    )

    _quadro_titulo_capa(
        doc,
        dados["periodicidade"],
        dados["tipo_ativo"],
        dados["subestacao"],
    )

    _quadro_corpo_tecnico(
        doc,
        corpo_tecnico,
        numero_os,
        numero_apr,
        dados["periodo"],
        concessao,
    )


# ============================================================
# NOVA SEÇÃO APÓS CAPA
# ============================================================

def _nova_secao_conteudo(
    doc: Document,
    titulo_cabecalho: str,
    subestacao: str,
):

    secao = doc.add_section(
        WD_SECTION.NEW_PAGE
    )

    secao.page_width = Cm(21)
    secao.page_height = Cm(29.7)

    secao.left_margin = Cm(1.8)
    secao.right_margin = Cm(1.8)

    secao.top_margin = Cm(1.6)
    secao.bottom_margin = Cm(1.5)

    # Não herdar capa
    secao.header.is_linked_to_previous = False
    secao.footer.is_linked_to_previous = False

    _iniciar_numeracao(
        secao,
        1
    )

    # ======================================
    # CABEÇALHO
    # ======================================

    header = secao.header

    p = header.paragraphs[0]

    p.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    p.paragraph_format.space_after = Pt(2)

    run = p.add_run(
        titulo_cabecalho.upper()
    )

    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.bold = False

    # ======================================
    # RODAPÉ
    # ======================================

    footer = secao.footer

    p = footer.paragraphs[0]

    p.alignment = (
        WD_ALIGN_PARAGRAPH.LEFT
    )

    r = p.add_run(
        "RIALMA TRANSMISSORA DE ENERGIA "
        f"V S.A. – SUBESTAÇÃO "
        f"{subestacao.upper()}"
    )

    r.font.name = "Arial"
    r.font.size = Pt(7)

    # Número da página à direita
    p2 = footer.add_paragraph()

    p2.alignment = (
        WD_ALIGN_PARAGRAPH.RIGHT
    )

    _campo_word(
        p2,
        "PAGE"
    )

    return secao


# ============================================================
# SUMÁRIO
# ============================================================

def _adicionar_sumario(doc: Document, grupos_fotos: list | None = None, tipo_ativo: str = "ATIVOS") -> None:
    """Gera um sumário estático, sem campos do Word ou LibreOffice."""
    grupos_fotos = grupos_fotos or []
    itens_por_pagina = 26
    quantidade_entradas = 5 + len(grupos_fotos)
    paginas_sumario = max(1, (quantidade_entradas + itens_por_pagina - 1) // itens_por_pagina)

    pagina_introducao = 2 + paginas_sumario
    pagina_inspecoes = pagina_introducao + 1
    pagina_ativo = pagina_inspecoes
    entradas = [
        ("1. INTRODUÇÃO", pagina_introducao, 0),
        ("2. PARÂMETROS NA INSPEÇÃO", pagina_introducao, 0),
        ("a. CONDIÇÕES DIVERSAS", pagina_introducao, 1),
        (f"3. INSPEÇÕES {tipo_ativo.upper()}", pagina_inspecoes, 0),
    ]

    total_paginas_fotos = 0
    for indice, (ativo, fotos_ativo) in enumerate(grupos_fotos, start=1):
        paginas_ativo = max(1, (len(fotos_ativo) + 3) // 4)
        entradas.append((f"3.{indice} - {ativo}", pagina_ativo, 1))
        pagina_ativo += paginas_ativo
        total_paginas_fotos += paginas_ativo

    pagina_anormalidades = pagina_inspecoes + max(1, total_paginas_fotos)
    entradas.append(("4. ANORMALIDADES ENCONTRADAS", pagina_anormalidades, 0))

    for indice_inicio in range(0, len(entradas), itens_por_pagina):
        if indice_inicio:
            doc.add_page_break()
        titulo = doc.add_paragraph()
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        titulo.paragraph_format.space_after = Pt(14)
        run = titulo.add_run("SUMÁRIO" if not indice_inicio else "SUMÁRIO (CONTINUAÇÃO)")
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(12)

        for texto, pagina, nivel in entradas[indice_inicio:indice_inicio + itens_por_pagina]:
            paragrafo = doc.add_paragraph()
            paragrafo.paragraph_format.left_indent = Cm(0.7 if nivel else 0)
            paragrafo.paragraph_format.space_before = Pt(0)
            paragrafo.paragraph_format.space_after = Pt(3)
            paragrafo.paragraph_format.tab_stops.add_tab_stop(Cm(16), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
            run_texto = paragrafo.add_run(texto)
            run_texto.font.name = "Arial"
            run_texto.font.size = Pt(9)
            run_texto.bold = nivel == 0
            run_pagina = paragrafo.add_run(f"\t{pagina}")
            run_pagina.font.name = "Arial"
            run_pagina.font.size = Pt(9)

    doc.add_page_break()

# ============================================================
# METADADOS
# ============================================================

def _metadados(
    doc: Document,
    dados: dict
) -> None:

    tabela = doc.add_table(
        rows=4,
        cols=4
    )

    tabela.style = "Table Grid"

    tabela.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    linhas = [
        (
            "Data",
            dados["data_br"],
            "Periodicidade",
            dados["periodicidade"],
        ),
        (
            "Subestação",
            dados["subestacao"],
            "Tipo de ativo",
            dados["tipo_ativo"],
        ),
        (
            "Arquivo fonte",
            dados["arquivo"],
            "Fotografias",
            str(
                dados["quantidade_fotos"]
            ),
        ),
        (
            "Responsável",
            dados["responsavel"],
            "Status",
            dados["status"],
        ),
    ]

    for i, linha in enumerate(linhas):

        for j, valor in enumerate(linha):

            _texto_celula(
                tabela.cell(i, j),
                str(valor),
                j % 2 == 0
            )


# ============================================================
# FOTOS
# ============================================================

def _fotos_zip(
    caminho_zip: Path
):

    with zipfile.ZipFile(
        caminho_zip
    ) as pacote:

        for entrada in pacote.infolist():

            if entrada.is_dir():
                continue

            extensao = Path(
                entrada.filename
            ).suffix.lower()

            if extensao not in EXTENSOES:
                continue

            yield (
                entrada.filename,
                pacote.read(entrada)
            )


def _agrupar_fotos_por_ativo(
    fotos: list[tuple[str, bytes]],
    revisoes: dict | None = None,
) -> list[tuple[str, list[tuple[str, bytes]]]]:
    revisoes = revisoes or {}
    validas = [
        (nome, conteudo)
        for nome, conteudo in fotos
        if revisoes.get(nome, {}).get("incluir", True)
    ]
    ordem_fases = {
        "AZ": 0, "AZUL": 0,
        "BR": 1, "BRANCA": 1, "BRANCO": 1,
        "VM": 2, "VERMELHA": 2, "VERMELHO": 2,
    }

    def chave_ordenacao(foto):
        revisao = revisoes.get(foto[0], {})
        fase = _normalizar(str(revisao.get("fase") or "")).replace(" ", "")
        return (
            _normalizar(str(revisao.get("ativo") or "")),
            _normalizar(str(revisao.get("item") or "")),
            ordem_fases.get(fase, 99),
            fase,
        )

    # Dentro de cada ativo: item do plano e depois AZ -> BR -> VM.
    validas.sort(key=chave_ordenacao)
    grupos: dict[str, list[tuple[str, bytes]]] = {}
    for nome, conteudo in validas:
        ativo = str(revisoes.get(nome, {}).get("ativo") or "ATIVO NÃO IDENTIFICADO").strip()
        grupos.setdefault(ativo, []).append((nome, conteudo))
    return list(grupos.items())


def _preencher_celula_foto(
    celula,
    conteudo: bytes,
    revisao: dict,
) -> None:
    celula.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    _set_cell_margins(celula, top=40, start=40, bottom=40, end=40)
    celula.text = ""
    p = celula.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    try:
        # Dimensões solicitadas: comprimento/largura 4,3 cm e altura 9,56 cm.
        p.add_run().add_picture(
            io.BytesIO(conteudo),
            width=Cm(4.3),
            height=Cm(9.56),
        )
    except Exception:
        p.add_run("Imagem não suportada")

    fase_legenda = f"Fase: {revisao.get('fase')}" if revisao.get("fase") else None
    partes = [revisao.get("item"), fase_legenda, revisao.get("valor"), revisao.get("status")]
    legenda_texto = " | ".join(str(parte) for parte in partes if parte)
    if legenda_texto:
        legenda = celula.add_paragraph(legenda_texto)
        legenda.alignment = WD_ALIGN_PARAGRAPH.CENTER
        legenda.paragraph_format.space_before = Pt(0)
        legenda.paragraph_format.space_after = Pt(0)
        for run in legenda.runs:
            run.font.name = "Arial"
            run.font.size = Pt(7.5)
            run.bold = True

    if revisao.get("observacao"):
        nota = celula.add_paragraph(str(revisao["observacao"]))
        nota.alignment = WD_ALIGN_PARAGRAPH.CENTER
        nota.paragraph_format.space_before = Pt(0)
        nota.paragraph_format.space_after = Pt(0)
        for run in nota.runs:
            run.font.name = "Arial"
            run.font.size = Pt(7)


def _grade_fotos(
    doc: Document,
    fotos: list[tuple[str, bytes]],
    numero_secao: int,
    tipo_ativo: str,
    revisoes: dict | None = None,
    grupos_fotos: list | None = None,
) -> None:
    revisoes = revisoes or {}
    grupos_fotos = grupos_fotos if grupos_fotos is not None else _agrupar_fotos_por_ativo(fotos, revisoes)

    doc.add_page_break()
    titulo = doc.add_paragraph(style="Heading 1")
    titulo.add_run(f"{numero_secao}. INSPEÇÕES {tipo_ativo.upper()}")

    if not grupos_fotos:
        doc.add_paragraph("Nenhuma evidência fotográfica disponível.")
        return

    primeira_grade = True
    for indice_ativo, (ativo, fotos_ativo) in enumerate(grupos_fotos, start=1):
        for inicio in range(0, len(fotos_ativo), 4):
            lote = fotos_ativo[inicio:inicio + 4]
            if not primeira_grade:
                doc.add_page_break()
            primeira_grade = False

            # Um único tópico por ativo; páginas adicionais mantêm somente as fotos/legendas.
            if inicio == 0:
                titulo_ativo = doc.add_paragraph(style="Heading 2")
                titulo_ativo.add_run(f"{numero_secao}.{indice_ativo} - {ativo}")

            tabela = doc.add_table(rows=2, cols=2)
            tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
            tabela.style = "Table Grid"
            tabela.autofit = False
            for linha in tabela.rows:
                tr_pr = linha._tr.get_or_add_trPr()
                cant_split = OxmlElement("w:cantSplit")
                tr_pr.append(cant_split)
                for celula in linha.cells:
                    celula.width = Cm(8.2)

            for posicao in range(4):
                celula = tabela.cell(posicao // 2, posicao % 2)
                if posicao < len(lote):
                    nome, conteudo = lote[posicao]
                    _preencher_celula_foto(celula, conteudo, revisoes.get(nome, {}))
                else:
                    celula.text = ""

# ============================================================
# PARÂMETROS DA INSPEÇÃO
# ============================================================

def _tabela_parametros_inspecao(
    doc: Document,
    parametros: dict | None = None,
):

    parametros = parametros or {}

    tabela = doc.add_table(
        rows=5,
        cols=3
    )

    tabela.style = "Table Grid"

    tabela.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    dados = [
        (
            "",
            "Inicio",
            "Fim",
        ),
        (
            "HORÁRIO",
            parametros.get(
                "hora_inicio",
                ""
            ),
            parametros.get(
                "hora_fim",
                ""
            ),
        ),
        (
            "TEMPERATURA AMBIENTE (°C)",
            parametros.get(
                "temperatura_inicio",
                ""
            ),
            parametros.get(
                "temperatura_fim",
                ""
            ),
        ),
        (
            "FREQUÊNCIA (Hz)",
            parametros.get(
                "frequencia_inicio",
                ""
            ),
            parametros.get(
                "frequencia_fim",
                ""
            ),
        ),
        (
            "TENSÃO BARRA (KV)",
            parametros.get(
                "tensao_inicio",
                ""
            ),
            parametros.get(
                "tensao_fim",
                ""
            ),
        ),
    ]

    for i, linha in enumerate(dados):

        for j, valor in enumerate(linha):

            celula = tabela.cell(
                i,
                j
            )

            celula.text = str(
                valor
            )

            celula.vertical_alignment = (
                WD_CELL_VERTICAL_ALIGNMENT.CENTER
            )

            p = celula.paragraphs[0]

            p.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
                if j > 0
                else WD_ALIGN_PARAGRAPH.LEFT
            )

            for run in p.runs:

                run.font.name = "Arial"
                run.font.size = Pt(9)

                if i == 0 or j == 0:
                    run.bold = True


# ============================================================
# RELATÓRIO
# ============================================================

def gerar_relatorio_word(
    relatorio,
    tipo_ativo: str,
    subestacao: str,
    responsavel: str,
    fotos_revisadas=None,

    # Novos dados da capa
    corpo_tecnico=None,
    numero_os="",
    numero_apr="",
    concessao="Rialma Transmissora de Energia - RTV",
    periodo_capa="",
    texto_introducao="",

    # Parâmetros opcionais
    parametros_inspecao=None,

    # Logo
    caminho_logo=CAMINHO_LOGO,
) -> Path:

    caminho_zip = Path(
        relatorio.caminho_arquivo
    )

    destino = caminho_zip.with_suffix(
        ".docx"
    )

    categoria = _categoria(
        tipo_ativo
    )

    fotos = list(
        _fotos_zip(
            caminho_zip
        )
    )

    # ========================================================
    # REVISÕES
    # ========================================================

    revisoes = {
        foto.nome_arquivo_zip: {

            "ativo":
                foto.ativo.codigo_ativo
                if foto.ativo
                else None,

            "item":
                foto.plano_item.nome_item
                if foto.plano_item
                else None,

            "fase":
                getattr(foto.ativo, "fase", None)
                if foto.ativo
                else None,
            "valor":
                foto.valor_medido,

            "status":
                foto.status_item,

            "observacao":
                foto.observacao,

            "incluir":
                foto.incluir,

        }

        for foto in (
            fotos_revisadas or []
        )
    }

    grupos_fotos = _agrupar_fotos_por_ativo(fotos, revisoes)

    data = relatorio.data_referencia

    periodicidade = (
        relatorio.periodicidade
        .replace("_", " ")
    )

    # Período da capa
    meses = [
        "",
        "Janeiro",
        "Fevereiro",
        "Março",
        "Abril",
        "Maio",
        "Junho",
        "Julho",
        "Agosto",
        "Setembro",
        "Outubro",
        "Novembro",
        "Dezembro",
    ]

    periodo = periodo_capa.strip() if periodo_capa else (
        f"{meses[data.month]}/ "
        f"{data.year}"
    )

    dados = {

        "data_br":
            data.strftime(
                "%d/%m/%Y"
            ),

        "periodo":
            periodo,

        "periodicidade":
            periodicidade,

        "subestacao":
            subestacao,

        "tipo_ativo":
            tipo_ativo,

        "arquivo":
            relatorio.nome_arquivo_original,

        "quantidade_fotos":
            len(fotos),

        "responsavel":
            responsavel,

        "status":
            relatorio.status,
    }

    # ========================================================
    # CORPO TÉCNICO
    # ========================================================

    if corpo_tecnico is None:

        corpo_tecnico = [
            {
                "nome": responsavel,
                "funcao": "",
            }
        ]

    # ========================================================
    # DOCUMENTO
    # ========================================================

    doc = Document()

    _configurar(doc)

    # ========================================================
    # CAPA
    # ========================================================

    _capa(
        doc=doc,
        dados=dados,
        corpo_tecnico=corpo_tecnico,
        numero_os=numero_os,
        numero_apr=numero_apr,
        concessao=concessao,
        caminho_logo=caminho_logo,
    )

    # ========================================================
    # SEÇÃO DO RELATÓRIO
    # ========================================================

    titulo_cabecalho = (
        f"RELATÓRIO DE MANUTENÇÃO "
        f"PREVENTIVA "
        f"{periodicidade.upper()} "
        f"DE {tipo_ativo.upper()}"
    )

    _nova_secao_conteudo(
        doc,
        titulo_cabecalho,
        subestacao,
    )

    # ========================================================
    # SUMÁRIO
    # ========================================================

    _adicionar_sumario(doc, grupos_fotos, tipo_ativo)

    # ========================================================
    # 1. INTRODUÇÃO
    # ========================================================

    h = doc.add_paragraph(
        style="Heading 1"
    )

    h.add_run(
        "1. INTRODUÇÃO"
    )

    introducao = texto_introducao.strip() if texto_introducao else (
        f"Este relatório reúne inspeções realizadas no dia {data.strftime('%d/%m/%Y')} na "
        f"Subestação {subestacao}, nos equipamentos do tipo {tipo_ativo}, operados pela "
        "Rialma Transmissora de Energia V S.A., referentes ao projeto da RTV.\n"
        "As inspeções foram realizadas tendo como referência os itens do plano de manutenção "
        "cadastrados no Sistema ENGVI.\n"
        "São apresentadas imagens dos equipamentos inspecionados. Os equipamentos que apresentarem "
        "anormalidades deverão ser destacados e vinculados às respectivas Solicitações de Serviço (SS)."
    )
    for paragrafo in introducao.splitlines():
        if paragrafo.strip():
            doc.add_paragraph(paragrafo.strip())

    # ========================================================
    # 2. PARÂMETROS
    # ========================================================

    h = doc.add_paragraph(
        style="Heading 1"
    )

    h.add_run(
        "2. PARAMETROS NA INSPEÇÃO"
    )

    h = doc.add_paragraph(
        style="Heading 2"
    )

    h.add_run(
        "a. CONDIÇÕES DIVERSAS"
    )

    _tabela_parametros_inspecao(
        doc,
        parametros_inspecao,
    )

    # ========================================================
    # 3. INSPEÇÕES
    # ========================================================

    _grade_fotos(
        doc=doc,
        fotos=fotos,
        numero_secao=3,
        tipo_ativo=tipo_ativo,
        revisoes=revisoes,
        grupos_fotos=grupos_fotos,
    )

    # ========================================================
    doc.add_page_break()

    # 4. ANORMALIDADES
    # ========================================================

    h = doc.add_paragraph(
        style="Heading 1"
    )

    h.add_run(
        "4. ANORMALIDADES ENCONTRADAS"
    )

    doc.add_paragraph(
        relatorio.observacao
        or
        "Durante as inspeções não foram encontradas "
        "novas anormalidades."
    )

    # ========================================================
    # PROPRIEDADES
    # ========================================================

    doc.core_properties.title = (
        f"RELATÓRIO "
        f"{periodicidade.upper()} "
        f"{tipo_ativo.upper()}"
    )

    doc.core_properties.subject = (
        "Relatório de manutenção preventiva "
        "com evidências fotográficas"
    )

    doc.core_properties.author = (
        "Sistema ENGVI"
    )

    # Faz o Word atualizar TOC/páginas ao abrir
    _atualizar_campos_ao_abrir(doc)

    # ========================================================
    # SALVAR
    # ========================================================

    doc.save(destino)

    return destino




